# -*- coding: utf-8 -*-
"""Generate a full product PRD for Banana Todo List (DOCX, Chinese)."""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ACCENT = RGBColor(0x00, 0xA6, 0x6E)        # green (BOSS直聘风)
ACCENT_DARK = RGBColor(0x0B, 0x7A, 0x5B)
GRAY = RGBColor(0x59, 0x59, 0x59)
BLACK = RGBColor(0x26, 0x26, 0x26)
BODY_FONT = "微软雅黑"
HDR_FILL = "E6F4EE"
OUT = r"F:\agents projects\AI_Note\backend\Banana_Todo_List_PRD.docx"


def set_east_asia(run, font=BODY_FONT):
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), font)
    rfonts.set(qn("w:hAnsi"), font)
    rfonts.set(qn("w:eastAsia"), font)


def set_doc_defaults(doc):
    section = doc.sections[0]
    section.page_width, section.page_height = Cm(21.0), Cm(29.7)
    section.top_margin, section.bottom_margin = Cm(2.0), Cm(2.0)
    section.left_margin, section.right_margin = Cm(2.2), Cm(2.2)
    style = doc.styles["Normal"]
    style.font.name = BODY_FONT
    style.font.size = Pt(10.5)
    style.font.color.rgb = BLACK
    style.element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT)
    pf = style.paragraph_format
    pf.space_after = Pt(4)
    pf.line_spacing = 1.3


def add_rule_bottom(paragraph, color="00A66E", size=10):
    ppr = paragraph._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), "3")
    bottom.set(qn("w:color"), color)
    pbdr.append(bottom)
    ppr.append(pbdr)


def shade_cell(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), fill)
    tcPr.append(shd)


def set_cell_size(cell, w_cm):
    tcPr = cell._tc.get_or_add_tcPr()
    tcW = OxmlElement("w:tcW")
    tcW.set(qn("w:w"), str(int(w_cm * 567)))
    tcW.set(qn("w:type"), "dxa")
    tcPr.append(tcW)


def h1(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run(text)
    r.font.name = BODY_FONT
    r.font.size = Pt(16)
    r.font.bold = True
    r.font.color.rgb = ACCENT_DARK
    set_east_asia(r)
    add_rule_bottom(p)
    return p


def h2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    r.font.name = BODY_FONT
    r.font.size = Pt(12.5)
    r.font.bold = True
    r.font.color.rgb = ACCENT_DARK
    set_east_asia(r)
    return p


def h3(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text)
    r.font.name = BODY_FONT
    r.font.size = Pt(11)
    r.font.bold = True
    r.font.color.rgb = BLACK
    set_east_asia(r)
    return p


def P(doc, text, indent=0, color=BLACK, size=10.5, bold=False):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.left_indent = Cm(indent)
    r = p.add_run(text)
    r.font.name = BODY_FONT
    r.font.size = Pt(size)
    r.font.bold = bold
    r.font.color.rgb = color
    set_east_asia(r)
    return p


def rich(doc, runs, indent=0):
    """runs: list of (text, kwargs)."""
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.left_indent = Cm(indent)
    for text, kw in runs:
        r = p.add_run(text)
        r.font.name = BODY_FONT
        r.font.size = Pt(kw.get("size", 10.5))
        r.font.bold = kw.get("bold", False)
        r.font.color.rgb = kw.get("color", BLACK)
        set_east_asia(r)
    return p


def bullet(doc, text, indent=0.4, marker="•", color=BLACK, bold_prefix=None):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(indent)
    p.paragraph_format.first_line_indent = Cm(-0.35)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(marker + " ")
    r.font.name = BODY_FONT
    r.font.size = Pt(10.5)
    r.font.color.rgb = ACCENT
    set_east_asia(r)
    if bold_prefix:
        r0 = p.add_run(bold_prefix)
        r0.font.name = BODY_FONT
        r0.font.size = Pt(10.5)
        r0.font.bold = True
        set_east_asia(r0)
    r2 = p.add_run(text)
    r2.font.name = BODY_FONT
    r2.font.size = Pt(10.5)
    r2.font.color.rgb = color
    set_east_asia(r2)
    return p


def code_block(doc, lines):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.left_indent = Cm(0.5)
    for i, line in enumerate(lines):
        if i:
            r = p.add_run("\n")
            set_east_asia(r, "Consolas")
        r = p.add_run(line)
        r.font.name = "Consolas"
        r.font.size = Pt(9.5)
        r.font.color.rgb = ACCENT_DARK
        rpr = r._element.get_or_add_rPr()
        rfonts = rpr.find(qn("w:rFonts"))
        if rfonts is None:
            rfonts = OxmlElement("w:rFonts")
            rpr.append(rfonts)
        rfonts.set(qn("w:ascii"), "Consolas")
        rfonts.set(qn("w:hAnsi"), "Consolas")
    return p


def data_table(doc, headers, rows, widths=None, font_size=9.5):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    for j, h in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(h)
        r.font.name = BODY_FONT
        r.font.size = Pt(font_size)
        r.font.bold = True
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_east_asia(r)
        shade_cell(cell, "00A66E")
    for row in rows:
        cells = table.add_row().cells
        for j, val in enumerate(row):
            cell = cells[j]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(str(val))
            r.font.name = BODY_FONT
            r.font.size = Pt(font_size)
            set_east_asia(r)
    if widths:
        for j, w in enumerate(widths):
            for row in table.rows:
                set_cell_size(row.cells[j], w)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def build():
    doc = Document()
    set_doc_defaults(doc)

    # ================= 封面 / 文档信息 =================
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(60)
    r = p.add_run("Banana Todo List")
    r.font.name = BODY_FONT
    r.font.size = Pt(30)
    r.font.bold = True
    r.font.color.rgb = ACCENT_DARK
    set_east_asia(r)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    r = p.add_run("产品需求文档（PRD）")
    r.font.name = BODY_FONT
    r.font.size = Pt(18)
    r.font.color.rgb = ACCENT
    set_east_asia(r)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(40)
    r = p.add_run("AI 原生待办助手 · Product Requirements Document")
    r.font.name = BODY_FONT
    r.font.size = Pt(11)
    r.font.color.rgb = GRAY
    set_east_asia(r)

    info = doc.add_table(rows=6, cols=2)
    info.alignment = WD_TABLE_ALIGNMENT.CENTER
    info.autofit = False
    rows = [
        ("文档名称", "Banana Todo List 产品需求文档（PRD）"),
        ("版本 / 状态", "V1.0 · 草案（Draft）"),
        ("作者", "Jiangshan2128"),
        ("编写日期", "2026-08-21"),
        ("适用范围", "AI Agent 后端服务 + 前端交互定义"),
        ("关联模块", "FastAPI 应用层 / LangGraph Agent 层 / 工具与记忆层"),
    ]
    for i, (label, value) in enumerate(rows):
        c0, c1 = info.rows[i].cells
        set_cell_size(c0, 3.2)
        set_cell_size(c1, 11.6)
        p0 = c0.paragraphs[0]
        r0 = p0.add_run(label)
        r0.font.name = BODY_FONT
        r0.font.size = Pt(10.5)
        r0.font.bold = True
        r0.font.color.rgb = GRAY
        set_east_asia(r0)
        p1 = c1.paragraphs[0]
        r1 = p1.add_run(value)
        r1.font.name = BODY_FONT
        r1.font.size = Pt(10.5)
        set_east_asia(r1)

    doc.add_page_break()

    # ================= 1. 产品概述 =================
    h1(doc, "1. 产品概述")

    h2(doc, "1.1 背景与问题")
    P(doc, "传统待办工具（备忘录、GTD / 清单类 App）的核心痛点在于“录入”与“维护”成本过高：")
    bullet(doc, "任务散落在对话、语音、笔记各处，需要手动转写与归类；")
    bullet(doc, "结构化字段（优先级、时间、依赖、周期）靠用户手工填写，多数用户懒得维护；")
    bullet(doc, "工具不理解用户偏好，不会主动记忆“我习惯怎么安排”；")
    bullet(doc, "写操作一旦发生即生效，误触/误记难以挽回。")
    P(doc, "随着大模型能力成熟，“自然语言对话即产品形态”成为可能。用户应当可以直接说一句"
       "“下周五之前把方案初稿做完，再提醒我准备周会材料”，产品就能自动完成任务的提取、拆解、"
       "排期与记录——而不是让用户去填一张表格。")

    h2(doc, "1.2 产品定位")
    P(doc, "Banana Todo List 是一款 **AI 原生个人待办助手**：以自然对话（文字 / 语音）作为唯一交互入口，"
       "由 LLM 智能体自动完成任务的提取、结构化、子任务拆解、时间排期与依赖建模，"
       "并通过**人机协同审批（HITL）**让用户对每一次写操作拥有最终控制权。"
       "产品以“会对话、会自我整理、会自动记忆偏好、先确认后写入”为核心体验，"
       "并可扩展接入钉钉办公生态与私有知识库问答。")

    h2(doc, "1.3 目标用户")
    data_table(doc,
        ["用户群体", "核心诉求", "典型行为"],
        [
            ["个人效率用户（白领 / 学生 / 创业者）", "任务多、时间紧，希望“说一句话就完成建任务”", "语音随手记、对话式排期、期望 AI 自动拆解"],
            ["钉钉办公用户", "待办需与钉钉待办联动、跨端同步", "授权钉钉账号，审批后同步待办"],
            ["知识工作者", "需要基于个人资料的问答与全局概览", "向知识库提问、获取多跳推理结论"],
        ],
        widths=[4.2, 6.0, 6.4])

    h2(doc, "1.4 产品价值主张")
    bullet(doc, "零成本录入：说话 / 打字即可建任务，AI 自动结构化。", bold_prefix="价值一 · ")
    bullet(doc, "智能整理：自动识别优先级、时间、依赖，支持子任务拆解与周期任务。", bold_prefix="价值二 · ")
    bullet(doc, "越用越懂你：长期记忆用户画像与计划偏好，逐次对话注入系统提示词。", bold_prefix="价值三 · ")
    bullet(doc, "安全可信：所有写操作先出提案、用户确认后落库（HITL），杜绝误操作。", bold_prefix="价值四 · ")
    bullet(doc, "生态打通：钉钉办公工具即插即用（MCP），私有知识库问答。", bold_prefix="价值五 · ")

    h2(doc, "1.5 竞品与差异化")
    data_table(doc,
        ["竞品类型", "代表", "局限", "本产品差异点"],
        [
            ["传统待办 App", "滴答清单 / Todoist / Microsoft To Do", "手动录入、字段维护成本高、无对话式交互", "对话 / 语音即录入 + AI 自动结构化 + HITL 确认"],
            ["通用 AI 助手", "ChatGPT 记忆 / 主流助手", "无任务专属数据模型、无审批流、不持久化业务记忆", "任务数据模型 + 审批流 + 长期业务记忆（PostgresStore）"],
            ["笔记 / 知识工具", "Notion AI", "偏文档场景，待办执行链路弱", "从“说”到“做”的闭环：建任务→排期→执行状态跟踪"],
        ],
        widths=[2.9, 3.4, 4.6, 5.7])

    # ================= 2. 名词解释 =================
    h1(doc, "2. 名词解释（术语表）")
    data_table(doc,
        ["术语", "说明"],
        [
            ["LangGraph", "基于图结构的智能体编排框架：状态图（StateGraph）、节点、条件边、checkpoint 持久化"],
            ["Agent（智能体）", "能够自主调用工具、维护多轮状态并完成目标任务的 LLM 应用"],
            ["HITL", "Human-in-the-Loop 人机协同：写操作先中断等待用户审批，再继续执行"],
            ["MCP", "Model Context Protocol，模型上下文协议：以统一方式接入外部工具（如钉钉）"],
            ["Tool（工具）", "Agent 可调用的外部能力（任务 CRUD、网页搜索、钉钉操作等）"],
            ["TrustCall", "确定性结构化抽取库：将 LLM 输出稳定映射到任务 / 画像等模型，返回可补丁的文档"],
            ["RAG", "Retrieval-Augmented Generation 检索增强生成：基于向量检索的知识库问答"],
            ["GraphRAG", "基于实体关系知识图谱的全局概览与多跳推理问答"],
            ["Checkpoint", "图运行状态快照：保存对话历史与中断点，支持断点续跑与 HITL 恢复"],
            ["SSE", "Server-Sent Events 服务端推送：用于 token 级流式输出"],
            ["Failover", "故障转移：主模型不可用时自动切换备用模型"],
        ],
        widths=[3.2, 13.4])

    # ================= 3. 产品架构 =================
    h1(doc, "3. 产品架构")

    h2(doc, "3.1 系统架构总览")
    P(doc, "系统为两层架构：**FastAPI Web 应用层**与 **LangGraph Agent 核心层**。"
       "Web 层在 lifespan 中通过依赖注入容器创建 store / graph / DingTalk runtime，"
       "并以 Depends 方式注入到请求处理器，与 Agent 层解耦。")
    code_block(doc, [
        "┌────────────────────────── FastAPI Web 层 ──────────────────────────┐",
        "│  auth  │  user  │  chat(文本/流式/恢复)  │  tasks  │  jobs  │  diag  │",
        "└────┬───────────────────────────────┬───────────────────────────────┘",
        "     │ 依赖注入容器 AppContext        │",
        "     ▼                               ▼",
        "┌────────────────── LangGraph Agent 层 ─────────────────────────────┐",
        "│  StateGraph: START → transcription? → agent → tools → HITL? → END │",
        "│  中间件流水线: MemoryLoad → SystemPrompt → ToolBinding → LLM       │",
        "│  工具层: 核心工具 + 按需推广 MCP 工具 + TrustCall 结构化抽取        │",
        "│  记忆层: PostgresStore(长期业务记忆) + MemorySaver(checkpoint)     │",
        "└───────────────────────────────────────────────────────────────────┘",
    ])

    h2(doc, "3.2 Agent 图拓扑")
    P(doc, "图采用**显式拓扑**而非简单的“LLM + 工具”循环，节点与条件边如下：")
    code_block(doc, [
        "  START ──▶ [transcription?] ──▶ agent ──▶ tools ──▶ [HITL?] ──▶ agent ──▶ END",
        "                                        ▲            │",
        "                                        └──── 循环 ────┘",
    ])
    bullet(doc, "transcription：语音输入走独立子图（私有 state、独立 checkpoint 命名空间），完成转写后汇入主图。")
    bullet(doc, "agent：中间件流水线组装系统提示词并调用模型（含工具绑定、多模型容错）。")
    bullet(doc, "tools：按用户隔离的 ScopedToolNode，执行核心工具与推广的 MCP 工具。")
    bullet(doc, "HITL：命中写操作时中断，等待用户审批后 resume 落库。")

    h2(doc, "3.3 分层说明")
    data_table(doc,
        ["层", "职责", "关键实现"],
        [
            ["Web 层 (app/)", "HTTP 入口、鉴权、参数解析、流式响应", "FastAPI + sse_starlette；DI 容器在 lifespan 创建"],
            ["图编排层 (agents/graph/)", "节点 / 路由 / 容错 / 状态定义", "StateGraph + conditional edges + set_node_defaults"],
            ["中间件层 (middleware/)", "横切关注点分离", "MemoryLoad → SystemPrompt → ToolBinding 俄套式流水线"],
            ["工具层 (tools/)", "Agent 可调能力", "8 个核心工具 + 按需推广的 MCP 工具 + TrustCall 抽取"],
            ["记忆层", "长期 / 短期记忆", "PostgresStore（业务数据）+ MemorySaver（对话与断点）"],
        ],
        widths=[3.4, 5.4, 7.8])

    # ================= 4. 核心用户场景 =================
    h1(doc, "4. 核心用户场景")
    for no, title, steps in [
        (1, "语音随手记任务", ["用户在首页按住麦克风说：“下周五交方案初稿”，", "系统先经转录子图转写为文本，再由 Agent 提取任务并生成提案，", "前端弹出审批卡片，用户点“确认”，任务落库。"]),
        (2, "对话拆解复杂任务", ["用户说：“帮我准备年会，事项列清楚”，", "Agent 自动拆解为子任务（场地 / 物料 / 议程 / 邀请），识别依赖关系（pre_task）与优先级，", "用户审批后一次性写入多任务。"]),
        (3, "多轮排期与记忆", ["用户多次提到“我习惯上午处理难事”，Agent 记入用户画像；", "下次排期时自动把高优先级任务排在上午，并在回复中说明依据。"]),
        (4, "任务执行与追踪", ["用户在对话中问“我这周还有哪些没做完”，Agent 检索任务列表按时间 / 优先级呈现，", "用户说“把 1 号标成已完成”，通过 HITL 确认后状态更新，并可同步钉钉待办。"]),
        (5, "知识库问答", ["用户向私有知识库提问“我们历次周会的结论有哪些”，", "RAG 检索相关片段作答；追问“相关项目之间是什么关系”时，GraphRAG 基于实体图谱做多跳推理。"]),
        (6, "异常恢复", ["主模型短暂不可用，系统自动降级备用模型，用户无感知；", "某节点超时后自动重试，最终仍失败时返回友好提示，不泄露内部异常。"]),
    ]:
        h3(doc, f"场景 {no} · {title}")
        for s in steps:
            bullet(doc, s)

    # ================= 5. 功能需求 =================
    h1(doc, "5. 功能需求")

    h2(doc, "5.0 功能需求总览")
    P(doc, "优先级说明：**P0** = MVP 必须（当前已实现）｜**P1** = 重要（下一迭代）｜**P2** = 增强 / 远期。")
    data_table(doc,
        ["功能模块", "P0（已实现）", "P1（进行中）", "P2（规划）"],
        [
            ["账号与鉴权", "Bearer Token 鉴权、用户数据隔离", "微信登录", "多端登录态同步"],
            ["对话交互", "文本对话、SSE 流式、多轮会话（thread）", "会话历史列表", "会话重命名 / 归档"],
            ["任务管理", "对话建任务、任务 CRUD API、优先级 / 标签 / 状态", "子任务拆解、前置依赖、周期任务", "任务看板 / 日历视图"],
            ["HITL 审批", "提案卡片、单项编辑 / 拒绝、resume 一致性", "钉钉同步二次确认", "批量审批"],
            ["智能记忆", "用户画像、计划偏好、任务记忆、运行时注入", "记忆主动回顾", "记忆编辑界面"],
            ["语音输入", "音频转写（Groq Whisper）", "多语言识别", "声纹 / 说话人分离"],
            ["知识库", "—", "RAG 问答、GraphRAG 全局概览", "文档上传管理"],
            ["钉钉集成", "按需工具加载能力", "授权绑定、待办同步、90+ 工具", "更多 MCP 生态"],
            ["异步任务", "—", "长任务 Job 创建 / 查询 / 恢复", "任务编排与重试面板"],
            ["可观测性", "/health、结构化日志、LangSmith", "/diag 诊断页", "错误监控告警"],
        ],
        widths=[2.6, 4.4, 5.0, 4.6])

    h2(doc, "5.1 账号与鉴权")
    bullet(doc, "请求携带 Supabase GoTrue 签发的 Bearer Token，服务端解析出 user_id 并用于所有数据隔离。", bold_prefix="[P0] ")
    bullet(doc, "无 Token 时降级为 default 用户（便于本地 / 测试）。", bold_prefix="[P0] ")
    bullet(doc, "支持微信登录（/auth/wechat-login）及 GoTrue 代理透传。", bold_prefix="[P1] ")

    h2(doc, "5.2 对话交互（核心）")
    bullet(doc, "文本消息处理：POST /chat，返回回复、最新任务列表及中断状态。", bold_prefix="[P0] ")
    bullet(doc, "SSE 流式：POST /chat/stream，基于 astream_events 实现 token 级流式输出，HITL 中断事件实时推送。", bold_prefix="[P0] ")
    bullet(doc, "多轮会话：前端生成 session_id 作为 LangGraph thread_id；同一 session 续聊，新 session 开新线程。", bold_prefix="[P0] ")
    bullet(doc, "鉴权隔离：不同 user_id 加载不同记忆命名空间。", bold_prefix="[P0] ")

    h2(doc, "5.3 任务管理（核心）")
    P(doc, "任务模型字段：title / description / tag(work|personal) / assignee / priority(P0|P1|P2) / "
       "time(日期，必填，默认今天) / pre_task(前置任务) / status(not started|in progress|done|archived) / recurrence(周期)。")
    bullet(doc, "对话生成任务：LLM 经 TrustCall 结构化抽取，自动解析相对日期（今天 / 明天 / 本周五 → YYYY-MM-DD）。", bold_prefix="[P0] ")
    bullet(doc, "任务 REST 接口：列表、单删、批量删、部分字段更新；删除 / 更新后可向对应会话注入通知，防止 LLM 误判重加。", bold_prefix="[P0] ")
    bullet(doc, "子任务拆解与依赖：拆分子任务并建立 pre_task 依赖关系。", bold_prefix="[P1] ")
    bullet(doc, "周期任务：支持 daily / weekly:mon,wed,fri 等周期规则。", bold_prefix="[P1] ")

    h2(doc, "5.4 人机协同审批 HITL（核心差异化）")
    P(doc, "目标：**任何写操作先提案、用户确认后落库**，从机制上杜绝 AI 误写。")
    bullet(doc, "LLM 调用写工具（update_tasks）→ TrustCall 产出提案 JSON，工具不直接写库。", bold_prefix="[P0] ")
    bullet(doc, "路由至 hitl_node → interrupt() 挂起图，SSE 推送审批卡片。", bold_prefix="[P0] ")
    bullet(doc, "用户操作：全部通过 / 单项拒绝（rejected_keys）/ 编辑后通过（edited_tasks）/ 全部拒绝。", bold_prefix="[P0] ")
    bullet(doc, "resume 从 checkpoint 恢复、不重新调用 LLM，保证结果一致；TrustCall 确定性 key 防止续跑错位。", bold_prefix="[P0] ")
    bullet(doc, "审批通过后可一键同步钉钉待办。", bold_prefix="[P1] ")

    h2(doc, "5.5 智能记忆")
    bullet(doc, "长期记忆：用户画像（姓名 / 职业 / 偏好）、计划偏好、任务数据存 PostgresStore，命名空间按 user_id 隔离。", bold_prefix="[P0] ")
    bullet(doc, "短期记忆：对话历史与图状态存 MemorySaver（checkpoint）。", bold_prefix="[P0] ")
    bullet(doc, "运行时注入：系统提示词每次从真实记忆动态拼装（SystemPromptMiddleware），并注入钉钉 union_id。", bold_prefix="[P0] ")
    bullet(doc, "无数据库时优雅降级为内存存储。", bold_prefix="[P0] ")

    h2(doc, "5.6 语音与音频转写")
    bullet(doc, "支持文本 + 音频同时传入；音频先经转录子图（Groq Whisper）转写，再汇入主 Agent 处理。", bold_prefix="[P0] ")
    bullet(doc, "转录子图私有 state + 独立 checkpoint 命名空间，不污染主图状态。", bold_prefix="[P0] ")
    bullet(doc, "支持指定音频语言（zh / en）。", bold_prefix="[P1] ")

    h2(doc, "5.7 知识库问答（RAG / GraphRAG）")
    bullet(doc, "RAG：Qdrant 向量库 + 稠密 / 稀疏混合检索（Qwen text-embedding-v4），支持片段级问答。", bold_prefix="[P1] ")
    bullet(doc, "GraphRAG：实体关系知识图谱，支持全局概览与多跳推理。", bold_prefix="[P1] ")
    bullet(doc, "文档上传与管理（待定入口）。", bold_prefix="[P2] ")

    h2(doc, "5.8 钉钉集成（MCP）")
    bullet(doc, "通过 MCP（stdio）接入钉钉办公套件 90+ 工具，配置驱动加载。", bold_prefix="[P1] ")
    bullet(doc, "三层工具绑定：核心工具常驻 + tool_search 常驻 + MCP 工具按需推广，解决模型约 55 个工具上限。", bold_prefix="[P0] ")
    bullet(doc, "按用户启用 / 停用（enable / disable），授权回调绑定 union_id。", bold_prefix="[P1] ")

    h2(doc, "5.9 异步任务（Job）")
    bullet(doc, "长耗时请求异步化：创建 / 查询 / 恢复（chat/jobs）。", bold_prefix="[P1] ")
    bullet(doc, "前端可轮询任务状态，完成后取回结果。", bold_prefix="[P1] ")

    h2(doc, "5.10 健康与诊断")
    bullet(doc, "/health：存活探针，校验 store 实际可用性并上报数据库后端类型。", bold_prefix="[P0] ")
    bullet(doc, "/diag/wechat：微信接入诊断。", bold_prefix="[P1] ")

    # ================= 6. 数据模型 =================
    h1(doc, "6. 数据模型")

    h2(doc, "6.1 任务 Task")
    data_table(doc,
        ["字段", "类型", "必填", "说明"],
        [
            ["title", "str", "是", "任务标题"],
            ["description", "str", "否", "任务详情"],
            ["tag", "work | personal", "否", "任务类别，默认 personal"],
            ["assignee", "str", "否", "责任人"],
            ["priority", "P0 | P1 | P2", "否", "P0=今日紧急，P1=重要，P2=常规；默认 P1"],
            ["time", "date", "是", "排期日期，默认今天；宽松解析多种日期格式"],
            ["pre_task", "str", "否", "前置任务标题（依赖建模）"],
            ["status", "not started | in progress | done | archived", "否", "执行状态，默认 not started"],
            ["recurrence", "str", "否", "周期规则：daily 或 weekly:mon,wed,fri"],
        ],
        widths=[2.6, 4.6, 1.4, 8.0])

    h2(doc, "6.2 用户画像 Profile")
    data_table(doc,
        ["字段", "类型", "说明"],
        [
            ["name / gender / age", "str / str / int", "基础信息"],
            ["job / location", "str", "职业与所在地"],
            ["description", "str", "用户偏好、计划习惯等自由描述（供提示词注入）"],
        ],
        widths=[4.0, 2.4, 10.2])

    h2(doc, "6.3 记忆隔离与命名空间")
    P(doc, "长期记忆以 (user_id) 作为命名空间前缀，短期记忆以 session_id（thread_id）作为线程键，"
       "保证多用户、多会话之间的数据完全隔离。")

    # ================= 7. 接口设计摘要 =================
    h1(doc, "7. 接口设计摘要")
    data_table(doc,
        ["方法", "路径", "说明"],
        [
            ["POST", "/api/v1/chat", "文本 / 音频对话（非流式）"],
            ["POST", "/api/v1/chat/stream", "SSE 流式对话，含 HITL 中断事件"],
            ["POST", "/api/v1/chat/resume", "携带用户审批决策继续被中断的图"],
            ["GET", "/api/v1/tasks/list", "获取当前用户全部任务"],
            ["DELETE", "/api/v1/tasks", "删除全部任务"],
            ["DELETE", "/api/v1/tasks/{key}", "删除单个任务"],
            ["PATCH", "/api/v1/tasks/{key}", "部分字段更新任务"],
            ["POST", "/api/v1/chat/jobs", "创建异步对话任务"],
            ["GET", "/api/v1/chat/jobs/{id}", "查询异步任务状态"],
            ["POST", "/api/v1/chat/jobs/{id}/resume", "恢复异步任务"],
            ["GET / PATCH", "/api/v1/user/profile", "获取 / 更新用户画像"],
            ["DELETE", "/api/v1/user", "注销账号"],
            ["POST", "/api/v1/auth/wechat-login", "微信登录"],
            ["GET", "/api/v1/dingtalk/status · authorize · callback", "钉钉授权绑定"],
            ["POST", "/api/v1/dingtalk/enable · disable", "启用 / 停用钉钉 MCP 工具"],
            ["GET", "/api/v1/diag/wechat", "微信接入诊断"],
            ["GET", "/health", "存活 / 数据库可用性探针"],
        ],
        widths=[2.6, 7.4, 6.6])

    # ================= 8. 非功能需求 =================
    h1(doc, "8. 非功能需求")

    h2(doc, "8.1 性能")
    bullet(doc, "SSE 首 token 延迟（TTFB）目标 < 1.5s；全量回复 P95 < 10s。")
    bullet(doc, "任务列表 / 常见 CRUD 接口 P95 < 200ms。")
    bullet(doc, "多模型 failover 切换开销控制在单次请求内（备用模型即时接管）。")

    h2(doc, "8.2 可用性与容错")
    bullet(doc, "图级容错：节点级重试（仅 429 / 5xx 与超时视为瞬时错误，4xx 致命）、按节点超时策略、统一错误兜底。")
    bullet(doc, "模型故障自动降级：FailoverChatModel 在 SDK 内建重试耗尽后切换备用供应商（如 DeepSeek → GLM）。")
    bullet(doc, "数据库不可用时自动降级内存存储，服务不中断。")
    bullet(doc, "HITL 中断绕过重试 / 超时，保证审批流程稳定。")

    h2(doc, "8.3 安全与隐私")
    bullet(doc, "全局异常处理绝不向客户端泄漏异常详情（DB 连接串、API Key、内部路径等）。")
    bullet(doc, "所有接口基于 Bearer Token 鉴权，数据按 user_id 命名空间强隔离。")
    bullet(doc, "第三方 MCP 工具 schema 入库前校验，避免畸形 schema 导致全请求失败。")

    h2(doc, "8.4 可观测性")
    bullet(doc, "LangSmith 全链路 trace（Agent 运行 / 工具调用 / LLM 请求）。")
    bullet(doc, "结构化日志 + 统一错误日志落点；/health 真实探活。")

    h2(doc, "8.5 兼容性与部署")
    bullet(doc, "部署于 CloudBase；MCP 服务按需拉起子进程，冷启动不预加载，保持启动快速。")
    bullet(doc, "Python 3.13；LangGraph >= 1.2（锁定 1.2.2）；支持 .env 配置化。")

    # ================= 9. 成功指标 =================
    h1(doc, "9. 成功指标")
    data_table(doc,
        ["指标", "定义", "目标"],
        [
            ["任务结构化成功率", "对话生成任务中一次通过结构化校验的比例", "≥ 90%"],
            ["HITL 审批通过率", "提案中用户直接通过（未编辑 / 未拒绝）的比例", "≥ 70%"],
            ["建任务转化率", "发起对话后成功落库任务的会话占比", "≥ 60%"],
            ["SSE 首 token 延迟", "从提交到首个 token 的耗时", "< 1.5s"],
            ["工具调用成功率", "Agent 发起的工具调用成功返回比例", "≥ 95%"],
            ["异常兜底体验", "图级错误被友好提示兜底的比例", "100%（不泄漏异常）"],
        ],
        widths=[4.0, 8.2, 4.4])

    # ================= 10. 迭代规划 =================
    h1(doc, "10. 迭代规划")
    data_table(doc,
        ["里程碑", "范围", "状态"],
        [
            ["M0 · MVP", "对话建任务、任务 CRUD、HITL 审批流、SSE 流式、语音转写、智能记忆、图级容错、LangSmith", "✅ 已完成"],
            ["M1 · 生态与深度", "多模型 failover、MCP 钉钉接入（90+ 工具）、RAG + GraphRAG 知识库、异步 Job、诊断页", "🔄 进行中"],
            ["M2 · 体验完善", "会话历史管理、周期任务提醒、任务看板 / 日历、前端体验打磨、移动端", "⏳ 规划中"],
            ["M3 · 远期", "多智能体协作、团队共享空间、技能市场（Agent 能力扩展）", "🔮 远期设想"],
        ],
        widths=[2.8, 9.6, 4.2])

    # ================= 11. 风险与对策 =================
    h1(doc, "11. 风险与对策")
    data_table(doc,
        ["风险", "等级", "对策"],
        [
            ["LLM 幻觉导致错误任务写入", "高", "写操作一律 HITL 审批；提示词注入权威性规则（本地任务为事实来源）；few-shot 反例约束"],
            ["模型不可用 / 不稳定", "高", "多模型配置 + FailoverChatModel 自动降级；图级重试（仅瞬时错误）"],
            ["工具数量超模型上限", "中", "三层工具绑定：核心常驻 + tool_search + 按需推广 MCP 工具"],
            ["多轮对话语义漂移", "中", "checkpoint 保存上下文；系统提示词实时注入当前记忆；相对日期解析避免“今天”永久错配"],
            ["用户数据隐私", "高", "按 user_id 命名空间隔离；Token 鉴权；异常不泄漏内部信息；第三方工具 schema 校验"],
            ["MCP 子进程资源开销", "中", "按需拉起、按用户启用 / 停用，避免常驻浪费"],
        ],
        widths=[5.6, 1.6, 9.4])

    # ================= 附：修订记录 =================
    h1(doc, "附录 · 修订记录")
    data_table(doc,
        ["版本", "日期", "修订内容", "作者"],
        [["V1.0", "2026-08-21", "初稿：基于当前后端实现产出完整产品 PRD", "Jiangshan2128"]],
        widths=[2.0, 3.0, 9.6, 3.0])

    doc.save(OUT)
    print(f"Saved: {OUT}")


if __name__ == "__main__":
    build()
