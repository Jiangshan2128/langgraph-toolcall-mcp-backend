# -*- coding: utf-8 -*-
"""Generate a professional AI Agent Developer resume template (DOCX, Chinese)."""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ---------- Palette ----------
ACCENT = RGBColor(0x1F, 0x4E, 0x79)      # deep blue
ACCENT_LIGHT = RGBColor(0x2E, 0x74, 0xB5)  # lighter blue
GRAY = RGBColor(0x59, 0x59, 0x59)
BLACK = RGBColor(0x26, 0x26, 0x26)
BODY_FONT = "微软雅黑"
OUT = r"F:\agents projects\AI_Note\backend\AI_Agent_开发简历模板.docx"


def set_east_asia(run, font=BODY_FONT):
    """Force the East-Asian font on a run so Chinese renders correctly."""
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), font)
    rfonts.set(qn("w:hAnsi"), font)
    rfonts.set(qn("w:eastAsia"), font)


def set_doc_defaults(doc):
    """A4 page, default font, narrow-ish margins."""
    section = doc.sections[0]
    section.page_width, section.page_height = Cm(21.0), Cm(29.7)
    section.top_margin = Cm(1.6)
    section.bottom_margin = Cm(1.6)
    section.left_margin = Cm(1.8)
    section.right_margin = Cm(1.8)

    style = doc.styles["Normal"]
    style.font.name = BODY_FONT
    style.font.size = Pt(10.5)
    style.font.color.rgb = BLACK
    style.element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT)
    pf = style.paragraph_format
    pf.space_after = Pt(2)
    pf.line_spacing = 1.25


def add_rule_bottom(paragraph, color="1F4E79", size=8):
    """Bottom border under a paragraph (used for section headings)."""
    ppr = paragraph._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), "2")
    bottom.set(qn("w:color"), color)
    pbdr.append(bottom)
    ppr.append(pbdr)


def add_heading(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.font.name = BODY_FONT
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.color.rgb = ACCENT
    set_east_asia(run)
    add_rule_bottom(p)
    return p


def add_bullet(doc, text, accent_word=None):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.25
    if accent_word and accent_word in text:
        idx = text.index(accent_word)
        pre, rest = text[:idx], text[idx:]
        run1 = p.add_run(pre)
        run1.font.name = BODY_FONT
        run1.font.size = Pt(10.5)
        set_east_asia(run1)
        run2 = p.add_run(rest)
        run2.font.name = BODY_FONT
        run2.font.size = Pt(10.5)
        run2.font.bold = True
        set_east_asia(run2)
    else:
        run = p.add_run(text)
        run.font.name = BODY_FONT
        run.font.size = Pt(10.5)
        set_east_asia(run)
    return p


def add_line(doc, left, right=None, bold_left=True, color_left=BLACK):
    """One line: bold left label + optional right/after text."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(left)
    run.font.name = BODY_FONT
    run.font.size = Pt(10.5)
    run.font.bold = bold_left
    run.font.color.rgb = color_left
    set_east_asia(run)
    if right:
        run2 = p.add_run("  " + right)
        run2.font.name = BODY_FONT
        run2.font.size = Pt(10.5)
        set_east_asia(run2)
    return p


def add_info_table(doc, rows):
    """Two-column info table used in the header."""
    table = doc.add_table(rows=len(rows), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = True
    for i, (label, value) in enumerate(rows):
        c0, c1 = table.rows[i].cells
        p0 = c0.paragraphs[0]
        p0.paragraph_format.space_after = Pt(2)
        r0 = p0.add_run(label)
        r0.font.name = BODY_FONT
        r0.font.size = Pt(10.5)
        r0.font.bold = True
        r0.font.color.rgb = GRAY
        set_east_asia(r0)
        p1 = c1.paragraphs[0]
        p1.paragraph_format.space_after = Pt(2)
        r1 = p1.add_run(value)
        r1.font.name = BODY_FONT
        r1.font.size = Pt(10.5)
        set_east_asia(r1)
    return table


def add_job_entry(doc, title, org, dates, bullets):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(f"{title}")
    r.font.name = BODY_FONT
    r.font.size = Pt(11)
    r.font.bold = True
    set_east_asia(r)
    r2 = p.add_run(f"　｜　{org}")
    r2.font.name = BODY_FONT
    r2.font.size = Pt(10.5)
    set_east_asia(r2)
    r3 = p.add_run(f"　{dates}")
    r3.font.name = BODY_FONT
    r3.font.size = Pt(10)
    r3.font.color.rgb = GRAY
    set_east_asia(r3)
    for b in bullets:
        add_bullet(doc, b)


def build():
    doc = Document()
    set_doc_defaults(doc)

    # ============ Header ============
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run("[您的姓名]")
    r.font.name = BODY_FONT
    r.font.size = Pt(26)
    r.font.bold = True
    r.font.color.rgb = ACCENT
    set_east_asia(r)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run("AI Agent 开发工程师")
    r.font.name = BODY_FONT
    r.font.size = Pt(13)
    r.font.color.rgb = ACCENT_LIGHT
    set_east_asia(r)

    add_info_table(doc, [
        ("电话：", "[138-0000-0000]"),
        ("邮箱：", "[yourname@example.com]"),
        ("所在地：", "[城市]"),
        ("GitHub / 博客：", "[github.com/yourname]"),
        ("最高学历：", "[本科 / 硕士 · 计算机相关专业]"),
        ("到岗时间：", "[随时 / 一个月内]"),
    ])

    # ============ 求职意向 ============
    add_heading(doc, "求职意向")
    add_line(doc, "目标职位：AI Agent 开发工程师 / LLM 应用工程师")
    add_line(doc, "期望薪资：", "[15K - 25K / 面议]")
    add_line(doc, "工作性质：", "[全职 / 实习]")

    # ============ 专业技能 ============
    add_heading(doc, "专业技能")
    for t in [
        "精通 LangGraph / LangChain 智能体编排，具备图状态机、条件路由、人机协同（Human-in-the-Loop）落地经验",
        "熟悉 MCP（Model Context Protocol）工具接入与调度，掌握 stdio 模式服务治理与工具按需绑定",
        "熟悉 RAG 检索增强生成与 GraphRAG 知识图谱，熟练使用 Qdrant + 向量模型（BGE / Qwen-Embedding）",
        "熟悉主流大模型 API（GLM / DeepSeek / OpenAI / Claude），掌握函数调用、结构化输出与流式响应",
        "掌握 FastAPI 服务开发与异步编程，了解 checkpoint 持久化、多轮会话恢复与容错重试策略",
        "熟悉 Python 全栈开发、Docker 部署与 CI/CD，具备生产环境问题排查经验",
    ]:
        add_bullet(doc, t)

    # ============ 工作经历 ============
    add_heading(doc, "工作经历")
    add_job_entry(
        doc,
        "[AI 应用开发工程师]",
        "[公司名称]",
        "[2023.07 - 至今]",
        [
            "负责 LLM 应用服务从 0 到 1 的架构设计与开发，主导 [核心模块] 落地，支撑 [N] 个业务方接入",
            "基于 LangGraph 构建任务型智能体，实现 HITL 审批流与断点续跑，任务处理准确率提升至 [XX]%",
            "搭建 RAG 知识库问答系统，优化检索排序策略，答案命中率提升 [XX]%",
            "推动 [工具/流程/规范] 改进，团队交付效率提升 [XX]%",
        ],
    )
    add_job_entry(
        doc,
        "[后端开发工程师]",
        "[公司名称]",
        "[2021.07 - 2023.06]",
        [
            "负责 [业务系统] 后端开发，支撑日活 [XX] 万，接口平均延迟 < [XX] ms",
            "设计并实现 [关键模块]，解决了 [具体难点]，获得 [成果/反馈]",
        ],
    )

    # ============ 项目经历 ============
    add_heading(doc, "项目经历")
    add_job_entry(
        doc,
        "Banana Todo List —— 任务型 AI Agent 后端（个人项目）",
        "LangGraph · FastAPI · MCP · Qdrant",
        "[2025.01 - 至今]",
        [
            "基于 LangGraph 构建“记忆 + 规划 + 工具调用”任务智能体，支持任务 CRUD、计划偏好记忆与多轮对话，通过 checkpoint 实现会话断点续跑",
            "设计 HITL 人机协同审批流：LLM 生成结构化任务提案（TrustCall）→ 用户审批 → 落库，杜绝误操作，提案键做确定性生成避免续跑时错位",
            "通过 MCP（stdio）接入钉钉办公套件 90+ 工具，采用“核心工具 + 按需推广”三级绑定策略，解决模型 55 个工具上限问题",
            "实现 RAG + GraphRAG 双引擎知识库：Qdrant 向量混合检索 + 实体关系图谱，支撑多跳推理与全局概览问答",
            "图级容错治理：自定义重试策略（仅 429/5xx 重试）、超时策略与统一错误兜底，生产可用性显著提升",
            "支持音频转写子图（Groq Whisper）与流式输出，构建了完整的智能体工程化闭环",
        ],
    )
    add_job_entry(
        doc,
        "[其他项目名称]",
        "[技术栈]",
        "[2023.01 - 2023.06]",
        [
            "[一句话介绍项目定位]",
            "[你的核心贡献与量化成果，尽量用数字说话]",
        ],
    )

    # ============ 教育背景 ============
    add_heading(doc, "教育背景")
    add_job_entry(
        doc,
        "[学校名称]",
        "[计算机科学与技术 · 本科]",
        "[2017.09 - 2021.06]",
        [
            "主修课程：[数据结构与算法 / 操作系统 / 计算机网络]",
            "[奖学金 / 竞赛获奖 / 相关活动]",
        ],
    )

    # ============ 自我评价 ============
    add_heading(doc, "自我评价")
    add_line(doc, "对 AI Agent 工程化有强烈热情，持续跟进 LangGraph / MCP 等最新技术演进；")
    add_line(doc, "具备从模型调用、工具编排到服务部署的全链路落地能力，注重代码质量与工程规范；")
    add_line(doc, "学习能力强、自我驱动，善于把新技术快速转化为可交付的业务价值。")

    doc.save(OUT)
    print(f"Saved: {OUT}")


if __name__ == "__main__":
    build()
