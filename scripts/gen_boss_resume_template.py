# -*- coding: utf-8 -*-
"""Generate a BOSS直聘-style resume template (DOCX, Chinese).

Replicates the look & layout of the BOSS直聘 online resume:
photo + name + tag chips header, 求职意向 block, education -> work -> projects.
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ---------- BOSS直聘-ish palette (green accent) ----------
ACCENT = RGBColor(0x00, 0xA6, 0x6E)          # teal green
ACCENT_DARK = RGBColor(0x0B, 0x7A, 0x5B)     # darker green text
ACCENT_BG = "E6F4EE"                          # light green chip bg
GRAY = RGBColor(0x59, 0x59, 0x59)
BLACK = RGBColor(0x26, 0x26, 0x26)
BODY_FONT = "微软雅黑"
OUT = r"F:\agents projects\AI_Note\backend\AI_Agent_开发简历模板_BOSS直聘版.docx"
TEXT_WIDTH_CM = 17.4  # A4 width 21 - left 1.8 - right 1.8


def set_east_asia(run, font=BODY_FONT):
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), font)
    rfonts.set(qn("w:hAnsi"), font)
    rfonts.set(qn("w:eastAsia"), font)


def set_doc_defaults(doc):
    section = doc.sections[0]
    section.page_width, section.page_height = Cm(21.0), Cm(29.7)
    section.top_margin = Cm(1.6)
    section.bottom_margin = Cm(1.6)
    section.left_margin = Cm(1.8)
    section.right_margin = Cm(1.8)
    style = doc.styles["Normal"]
    style.font.name = BODY_FONT
    style.font.size = Pt(10.5)
    style.font.color.rgb = BLACK
    style.element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT)
    pf = style.paragraph_format
    pf.space_after = Pt(2)
    pf.line_spacing = 1.25


def add_rule_bottom(paragraph, color="00A66E", size=8):
    ppr = paragraph._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), "2")
    bottom.set(qn("w:color"), color)
    pbdr.append(bottom)
    ppr.append(pbdr)


def shade_cell(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), fill)
    tcPr.append(shd)


def set_cell_size(cell, w_cm, h_cm=None):
    tcPr = cell._tc.get_or_add_tcPr()
    tcW = OxmlElement("w:tcW")
    tcW.set(qn("w:w"), str(int(w_cm * 567)))
    tcW.set(qn("w:type"), "dxa")
    tcPr.append(tcW)
    if h_cm:
        tcH = OxmlElement("w:trHeight")
        tcH.set(qn("w:val"), str(int(h_cm * 567)))
        tcH.set(qn("w:hRule"), "atLeast")
        tcPr.append(tcH)


def add_dashed_border(cell, color="999999"):
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        e = OxmlElement(f"w:{edge}")
        e.set(qn("w:val"), "dashed")
        e.set(qn("w:sz"), "8")
        e.set(qn("w:color"), color)
        tcBorders.append(e)
    tcPr.append(tcBorders)


def add_heading(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    r0 = p.add_run("■ ")
    r0.font.color.rgb = ACCENT
    r0.font.size = Pt(12)
    r0.font.bold = True
    set_east_asia(r0)
    r = p.add_run(text)
    r.font.bold = True
    r.font.size = Pt(13)
    r.font.color.rgb = BLACK
    set_east_asia(r)
    add_rule_bottom(p, color="00A66E")
    return p


def add_tag_chips(doc, tags):
    """A row of BOSS直聘-style rounded tag chips (shaded table cells)."""
    table = doc.add_table(rows=1, cols=len(tags))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    for i, tag in enumerate(tags):
        cell = table.rows[0].cells[i]
        set_cell_size(cell, len(tag) * 0.62 + 0.7, 0.62)
        shade_cell(cell, ACCENT_BG)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(tag)
        r.font.size = Pt(9.5)
        r.font.color.rgb = ACCENT_DARK
        set_east_asia(r)


def add_para(doc, runs, space_after=2, align=None):
    """Add a paragraph from a list of (text, kwargs) run specs."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    if align:
        p.alignment = align
    for text, kw in runs:
        r = p.add_run(text)
        r.font.name = BODY_FONT
        r.font.size = kw.get("size", Pt(10.5))
        r.font.bold = kw.get("bold", False)
        r.font.color.rgb = kw.get("color", BLACK)
        set_east_asia(r)
    return p


def add_dated_line(doc, left_runs, date, tab_at=TEXT_WIDTH_CM):
    """Left content + right-aligned date using a right tab stop."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.tab_stops.add_tab_stop(Cm(tab_at), WD_TAB_ALIGNMENT.RIGHT)
    for text, kw in left_runs:
        r = p.add_run(text)
        r.font.name = BODY_FONT
        r.font.size = kw.get("size", Pt(10.5))
        r.font.bold = kw.get("bold", False)
        r.font.color.rgb = kw.get("color", BLACK)
        set_east_asia(r)
    r = p.add_run("\t" + date)
    r.font.name = BODY_FONT
    r.font.size = Pt(9.5)
    r.font.color.rgb = GRAY
    set_east_asia(r)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.25
    r = p.add_run("· ")
    r.font.name = BODY_FONT
    r.font.size = Pt(10.5)
    r.font.color.rgb = ACCENT
    set_east_asia(r)
    r2 = p.add_run(text)
    r2.font.name = BODY_FONT
    r2.font.size = Pt(10.5)
    set_east_asia(r2)
    return p


def add_entry(doc, title_runs, date, bullets):
    add_dated_line(doc, title_runs, date)
    for b in bullets:
        add_bullet(doc, b)


def build():
    doc = Document()
    set_doc_defaults(doc)

    # ============ Header: photo | name + chips ============
    header = doc.add_table(rows=1, cols=2)
    header.alignment = WD_TABLE_ALIGNMENT.LEFT
    header.autofit = False

    # -- photo placeholder cell --
    pc = header.rows[0].cells[0]
    set_cell_size(pc, 3.2, 3.8)
    add_dashed_border(pc)
    pc.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    pp = pc.paragraphs[0]
    pp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pr = pp.add_run("照片")
    pr.font.size = Pt(12)
    pr.font.color.rgb = GRAY
    set_east_asia(pr)

    # -- right info cell --
    ic = header.rows[0].cells[1]
    ic.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    p = ic.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run("[您的姓名]")
    r.font.name = BODY_FONT
    r.font.size = Pt(22)
    r.font.bold = True
    r.font.color.rgb = BLACK
    set_east_asia(r)

    p = ic.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run("AI Agent 开发工程师")
    r.font.name = BODY_FONT
    r.font.size = Pt(12)
    r.font.color.rgb = ACCENT_DARK
    r.font.bold = True
    set_east_asia(r)

    add_tag_chips(ic, ["3年经验", "本科", "北京"])
    add_para(ic, [
        ("电话：", {"size": Pt(10), "color": GRAY}),
        ("[138-0000-0000]", {"size": Pt(10)}),
        ("　　邮箱：", {"size": Pt(10), "color": GRAY}),
        ("[yourname@example.com]", {"size": Pt(10)}),
    ], space_after=0)

    # ============ 求职意向 ============
    add_heading(doc, "求职意向")
    intent = doc.add_table(rows=3, cols=4)
    intent.alignment = WD_TABLE_ALIGNMENT.LEFT
    intent.autofit = True
    pairs = [
        ("期望职位", "AI Agent 开发工程师"),
        ("期望薪资", "[15K - 25K]"),
        ("工作性质", "[全职]"),
        ("期望城市", "[北京]"),
        ("到岗时间", "[随时]"),
        ("工作状态", "[在职-看机会]"),
    ]
    for i, (label, value) in enumerate(pairs):
        r_, c_ = divmod(i, 2)
        cell = intent.rows[r_].cells[c_ * 2]
        cell.paragraphs[0].paragraph_format.space_after = Pt(2)
        lr = cell.paragraphs[0].add_run(label + "：")
        lr.font.name = BODY_FONT
        lr.font.size = Pt(10.5)
        lr.font.color.rgb = GRAY
        lr.font.bold = True
        set_east_asia(lr)
        vr = cell.paragraphs[0].add_run(value)
        vr.font.name = BODY_FONT
        vr.font.size = Pt(10.5)
        set_east_asia(vr)

    # ============ 教育经历 ============
    add_heading(doc, "教育经历")
    add_entry(
        doc,
        [("[学校名称]", {"bold": True}), ("　[计算机科学与技术] · 本科", {}), ("　[GPA 3.5/4.0]", {"color": GRAY, "size": Pt(9.5)})],
        "[2017.09 - 2021.06]",
        ["主修课程：数据结构与算法 / 操作系统 / 计算机网络",
         "[奖学金 / 竞赛获奖 / 相关经历]"],
    )

    # ============ 工作经历 ============
    add_heading(doc, "工作经历")
    add_entry(
        doc,
        [("[公司名称]", {"bold": True}), ("　[AI 应用开发工程师]", {})],
        "[2023.07 - 至今]",
        ["负责 LLM 应用服务从 0 到 1 的架构设计与开发，主导 [核心模块] 落地，支撑 [N] 个业务方接入",
         "基于 LangGraph 构建任务型智能体，实现 HITL 审批流与断点续跑，任务处理准确率提升至 [XX]%",
         "搭建 RAG 知识库问答系统，优化检索排序策略，答案命中率提升 [XX]%",
         "推动 [工具/流程/规范] 改进，团队交付效率提升 [XX]%"],
    )
    add_entry(
        doc,
        [("[公司名称]", {"bold": True}), ("　[后端开发工程师]", {})],
        "[2021.07 - 2023.06]",
        ["负责 [业务系统] 后端开发，支撑日活 [XX] 万，接口平均延迟 < [XX] ms",
         "设计并实现 [关键模块]，解决了 [具体难点]，获得 [成果/反馈]"],
    )

    # ============ 项目经历 ============
    add_heading(doc, "项目经历")
    add_entry(
        doc,
        [("Banana Todo List —— 任务型 AI Agent 后端（个人项目）", {"bold": True})],
        "[2025.01 - 至今]",
        ["基于 LangGraph 构建“记忆 + 规划 + 工具调用”任务智能体，支持任务 CRUD、计划偏好记忆与多轮对话，通过 checkpoint 实现会话断点续跑",
         "设计 HITL 人机协同审批流：LLM 生成结构化任务提案（TrustCall）→ 用户审批 → 落库，杜绝误操作，提案键做确定性生成避免续跑时错位",
         "通过 MCP（stdio）接入钉钉办公套件 90+ 工具，采用“核心工具 + 按需推广”三级绑定策略，解决模型 55 个工具上限问题",
         "实现 RAG + GraphRAG 双引擎知识库：Qdrant 向量混合检索 + 实体关系图谱，支撑多跳推理与全局概览问答",
         "图级容错治理：自定义重试策略（仅 429/5xx 重试）、超时策略与统一错误兜底，生产可用性显著提升",
         "支持音频转写子图（Groq Whisper）与流式输出，构建了完整的智能体工程化闭环"],
    )
    add_entry(
        doc,
        [("[其他项目名称]", {"bold": True}), ("　[我的角色]", {})],
        "[2023.01 - 2023.06]",
        ["[一句话介绍项目定位]",
         "[你的核心贡献与量化成果，尽量用数字说话]"],
    )

    # ============ 专业技能 ============
    add_heading(doc, "专业技能")
    for t in [
        "精通 LangGraph / LangChain 智能体编排，具备图状态机、条件路由、人机协同（Human-in-the-Loop）落地经验",
        "熟悉 MCP（Model Context Protocol）工具接入与调度，掌握 stdio 模式服务治理与工具按需绑定",
        "熟悉 RAG 检索增强生成与 GraphRAG 知识图谱，熟练使用 Qdrant + 向量模型（BGE / Qwen-Embedding）",
        "熟悉主流大模型 API（GLM / DeepSeek / OpenAI / Claude），掌握函数调用、结构化输出与流式响应",
        "掌握 FastAPI 服务开发与异步编程，了解 checkpoint 持久化、多轮会话恢复与容错重试策略",
        "熟悉 Python 全栈开发、Docker 部署与 CI/CD，具备生产环境问题排查经验",
    ]:
        add_bullet(doc, t)

    # ============ 自我评价 ============
    add_heading(doc, "自我评价")
    for t in [
        "对 AI Agent 工程化有强烈热情，持续跟进 LangGraph / MCP 等最新技术演进；",
        "具备从模型调用、工具编排到服务部署的全链路落地能力，注重代码质量与工程规范；",
        "学习能力强、自我驱动，善于把新技术快速转化为可交付的业务价值。",
    ]:
        add_bullet(doc, t)

    doc.save(OUT)
    print(f"Saved: {OUT}")


if __name__ == "__main__":
    build()
